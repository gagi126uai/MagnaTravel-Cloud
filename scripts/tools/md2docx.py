#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
md2docx.py — Conversor simple de Markdown a Word (.docx) para los documentos
de MagnaTravel-Cloud. Pensado para que Gaston pueda LEER la documentacion en
Word sin tener que abrir los .md del repositorio.

Soporta el subconjunto de Markdown que usamos en docs/: titulos (#, ##, ###),
negrita (**texto**), `codigo inline`, listas con - y 1., tablas | a | b |,
citas con >, bloques de codigo ```...``` y separadores ---.

Uso:
    python scripts/tools/md2docx.py entrada.md salida.docx

No pretende ser un parser perfecto de Markdown; cubre lo que escribimos nosotros.
"""
import re
import sys

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_runs_with_inline(paragraph, text):
    """Agrega 'text' a un parrafo interpretando **negrita** y `codigo inline`."""
    # Partimos por los marcadores manteniendo los delimitadores.
    tokens = re.split(r"(\*\*.+?\*\*|`.+?`)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        else:
            paragraph.add_run(tok)


def is_table_row(line):
    return line.strip().startswith("|") and line.strip().endswith("|")


def is_table_separator(line):
    # | --- | --- | (la fila que separa el header del cuerpo)
    return bool(re.match(r"^\s*\|?[\s:|-]+\|?\s*$", line)) and "-" in line


def parse_table_row(line):
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def convert(md_path, docx_path):
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()
    # Estilo base un poco mas legible.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Bloque de codigo ```
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # saltar el ``` de cierre
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            continue

        # Tabla
        if is_table_row(line) and (i + 1 < n) and is_table_separator(lines[i + 1]):
            header = parse_table_row(line)
            i += 2  # saltar header + separador
            body = []
            while i < n and is_table_row(lines[i]):
                body.append(parse_table_row(lines[i]))
                i += 1
            cols = len(header)
            table = doc.add_table(rows=1, cols=cols)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            for c in range(cols):
                hdr[c].paragraphs[0].add_run(
                    header[c] if c < len(header) else ""
                ).bold = True
            for row in body:
                cells = table.add_row().cells
                for c in range(cols):
                    val = row[c] if c < len(row) else ""
                    add_runs_with_inline(cells[c].paragraphs[0], val)
            doc.add_paragraph()
            continue

        # Separador ---
        if stripped == "---":
            doc.add_paragraph().add_run("_" * 40)
            i += 1
            continue

        # Titulos
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
            i += 1
            continue

        # Cita >
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            run_text = stripped.lstrip(">").strip()
            r = p.add_run("» ")
            r.italic = True
            add_runs_with_inline(p, run_text)
            for rn in p.runs:
                rn.italic = True
            i += 1
            continue

        # Lista numerada
        m_num = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if m_num:
            p = doc.add_paragraph(style="List Number")
            add_runs_with_inline(p, m_num.group(3))
            i += 1
            continue

        # Lista con viñeta
        m_bul = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m_bul:
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_inline(p, m_bul.group(2))
            i += 1
            continue

        # Linea vacia
        if stripped == "":
            i += 1
            continue

        # Parrafo normal
        p = doc.add_paragraph()
        add_runs_with_inline(p, stripped)
        i += 1

    doc.save(docx_path)
    print(f"OK: {md_path} -> {docx_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Uso: python scripts/tools/md2docx.py entrada.md salida.docx")
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
